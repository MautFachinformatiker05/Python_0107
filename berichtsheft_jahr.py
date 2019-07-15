from docx import Document
import datetime
# from datetime import datetime, timedelta

import os

jahr, none, none= datetime.date.today().isocalendar()

for woche in range(1,54):
    document = Document('heft.docx')

    montag = datetime.datetime.strptime((str(jahr)+"-W"+str(woche-1)+"-1"), "%Y-W%W-%w").date()

    dateiname = "jahr\\" + "KW"+str(woche)+" "
    dateiname += str(montag)
    dateiname += " bis "
    dateiname += str(montag + datetime.timedelta(days=4))
    dateiname += "_Filip Golanski.docx"

    daten = ["D0", "D1", "D2", "D3", "D4"]

    for n in range(5):
            daten[n] = ((montag + datetime.timedelta(days=n)).strftime("%d.%m.%Y"))

    for zelle in document.tables[0]._cells:
        if zelle.text == "DATUM0":
            zelle.text = daten[0]
        if zelle.text == "DATUM1":
            zelle.text = daten[1]
        if zelle.text == "DATUM2":
            zelle.text = daten[2]
        if zelle.text == "DATUM3":
            zelle.text = daten[3]
        if zelle.text == "DATUM4":
            zelle.text = daten[4]

    for zelle in document.tables[1]._cells:
        zelle.text = zelle.text.replace("DATUM3",daten[3])

    print(dateiname + "\n")
    document.save(dateiname)

